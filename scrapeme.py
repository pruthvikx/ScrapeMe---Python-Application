import requests
from bs4 import BeautifulSoup
from googlesearch import search
from docx import Document
import sublist3r
import socket
from vulners import Vulners

# Function to perform Google search
def google_search(query, num_results):
    try:
        return list(search(query, stop=num_results, pause=2))
    except Exception as e:
        print(f"An error occurred while searching: {e}")
        return []

# Function to gather subdomains
def get_subdomains(domain):
    subdomains = sublist3r.main(domain, 40, savefile=None, ports=None, silent=True, verbose=False, enable_bruteforce=False, engines=None)
    return subdomains

# Function to perform port scanning
def port_scan(domain):
    common_ports = [21, 22, 23, 25, 53, 80, 110, 143, 443, 3389]
    open_ports = []
    for port in common_ports:
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        socket.setdefaulttimeout(1)
        result = sock.connect_ex((domain, port))
        if result == 0:
            open_ports.append(port)
        sock.close()
    return open_ports

# Function to gather technical information
def get_technical_info(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.title.string if soup.title else 'No title'
        headers = response.headers
        return {'title': title, 'headers': headers}
    except requests.exceptions.RequestException as e:
        print(f"Failed to retrieve {url}: {e}")
        return None

# Function to perform vulnerability assessment
def vulnerability_assessment(technical_info):
    vulners_api = Vulners(api_key='ENTER_YOUR_OWN_BRO')
    headers = technical_info.get('headers', {})
    vulnerabilities = []

    # Simple check against server header
    server = headers.get('Server', None)
    if server:
        try:
            vulns = vulners_api.find_all(f'software:"{server}"')
            for vuln in vulns:
                vulnerabilities.append(vuln.get('description', 'No description available'))
        except Exception as e:
            print(f"Failed to search vulnerabilities: {e}")

    return vulnerabilities

# Function to save results to document
def save_to_document(results, filename):
    doc = Document()
    for result in results:
        doc.add_heading(result['url'], level=1)
        doc.add_heading(result['data']['title'], level=2)
        for paragraph in result['data']['paragraphs']:
            doc.add_paragraph(paragraph)
        doc.add_heading("Subdomains", level=3)
        for subdomain in result['subdomains']:
            doc.add_paragraph(subdomain)
        doc.add_heading("Open Ports", level=3)
        for port in result['open_ports']:
            doc.add_paragraph(str(port))
        doc.add_heading("Technical Information", level=3)
        for key, value in result['technical_info'].items():
            doc.add_paragraph(f"{key}: {value}")
        doc.add_heading("Vulnerabilities", level=3)
        for vuln in result['vulnerabilities']:
            doc.add_paragraph(vuln)
    doc.save(filename)

# Main function
def main(query, num_results, output_filename):
    urls = google_search(query, num_results)
    results = []
    for url in urls:
        domain = url.split("//")[-1].split("/")[0]
        subdomains = get_subdomains(domain)
        open_ports = port_scan(domain)
        technical_info = get_technical_info(url)
        vulnerabilities = vulnerability_assessment(technical_info) if technical_info else []
        data = extract_data(scrape_url(url))
        results.append({
            'url': url,
            'data': data,
            'subdomains': subdomains,
            'open_ports': open_ports,
            'technical_info': technical_info,
            'vulnerabilities': vulnerabilities
        })
    save_to_document(results, output_filename)

# Function to scrape the URL
def scrape_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        return soup
    except requests.exceptions.RequestException as e:
        print(f"Failed to retrieve {url}: {e}")
        return None

# Function to extract data from the soup object
def extract_data(soup):
    data = {}
    title = soup.title.string if soup.title else 'No title'
    data['title'] = title
    paragraphs = [p.get_text() for p in soup.find_all('p')]
    data['paragraphs'] = paragraphs
    return data

if __name__ == "__main__":
    query = "https://www.geeksforgeeks.org/"
    num_results = 5
    output_filename = "geeknew.docx"
    main(query, num_results, output_filename)
